import os
import re
import streamlit as st
from typing import List, Dict
from io import BytesIO
from docx import Document
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI
from langchain.chains.summarize import load_summarize_chain
from langchain.docstore.document import Document as LangchainDocument
from crewai import Agent, Task, Crew, Process
import sys

class StreamToExpander:
    def __init__(self, expander, buffer_limit=10000):
        self.expander = expander
        self.buffer = []
        self.buffer_limit = buffer_limit

    def write(self, data):
        cleaned_data = re.sub(r'\x1B\[\d+;?\d*m', '', data)
        if len(self.buffer) >= self.buffer_limit:
            self.buffer.pop(0)
        self.buffer.append(cleaned_data)

        if "\n" in data:
            self.expander.markdown(''.join(self.buffer), unsafe_allow_html=True)
            self.buffer.clear()

    def flush(self):
        if self.buffer:
            self.expander.markdown(''.join(self.buffer), unsafe_allow_html=True)
            self.buffer.clear()

class DocumentProcessor:
    def __init__(self, api_key: str):
        self.embeddings = OpenAIEmbeddings(openai_api_key=api_key)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", ""]
        )
        self.llm = ChatOpenAI(
            temperature=0.2,
            model="gpt-4o-mini",
            openai_api_key=api_key
        )
        self.summarize_chain = load_summarize_chain(
            self.llm,
            chain_type="map_reduce",
            verbose=True
        )

    def extract_text_from_docx(self, file_bytes) -> str:
        doc = Document(BytesIO(file_bytes))
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    def create_document_embeddings(self, documents: List[Dict[str, str]]) -> FAISS:
        """Create embeddings for documents and store in FAISS index"""
        texts = []
        metadatas = []
        
        for doc in documents:
            splits = self.text_splitter.split_text(doc['content'])
            texts.extend(splits)
            metadatas.extend([{'source': doc['title']} for _ in splits])
        
        return FAISS.from_texts(texts, self.embeddings, metadatas=metadatas)

    def extract_tasks(self, rfp_content: str) -> List[str]:
        """Extract tasks from RFP using LLM"""
        prompt = f"""
        Extract the specific tasks or requirements from the following RFP content.
        Format each task as a separate item.
        
        RFP Content:
        {rfp_content}
        
        Tasks:"""
        
        response = self.llm.invoke(prompt)
        tasks = [task.strip() for task in response.content.split('\n') if task.strip()]
        return tasks

    def find_relevant_experience(self, task: str, vector_store: FAISS, top_k: int = 3) -> List[Dict]:
        """Find relevant past experience for a specific task"""
        results = vector_store.similarity_search_with_score(task, k=top_k)
        return [
            {
                'content': result[0].page_content,
                'source': result[0].metadata['source'],
                'relevance': result[1]
            }
            for result in results
        ]

    def summarize_experience(self, experiences: List[Dict]) -> str:
        """Summarize relevant past experiences into a cohesive narrative"""
        if not experiences:
            return ""
        
        combined_text = "\n\n".join([exp['content'] for exp in experiences])
        docs = [LangchainDocument(page_content=combined_text)]
        
        summary = self.summarize_chain.run(docs)
        return summary

class ProposalWriter:
    def __init__(self, api_key: str):
        self.doc_processor = DocumentProcessor(api_key)
        self.llm = ChatOpenAI(
            temperature=0.2,
            model="gpt-4o-mini",
            openai_api_key=api_key
        )

    def setup_agents(self):
        analyzer = Agent(
            role='Document Analyzer',
            goal='Analyze RFP requirements and find matching past experience',
            backstory="""Expert at analyzing documents and matching past performance to requirements.
                        Uses advanced NLP to identify relevant experience.""",
            llm=self.llm
        )

        writer = Agent(
            role='Proposal Writer',
            goal='Create compelling proposal sections with past performance evidence',
            backstory="""Expert proposal writer who creates persuasive content backed by
                        specific past performance examples.""",
            llm=self.llm
        )

        editor = Agent(
            role='Proposal Editor',
            goal='Polish and enhance proposal content',
            backstory="""Senior editor who ensures proposal quality and strong
                        past performance narratives.""",
            llm=self.llm
        )

        return analyzer, writer, editor

    def create_tasks(self, analyzer, writer, editor, rfp_tasks: List[str], 
                    past_experience: Dict[str, List[Dict]]):
        tasks = []
        
        # Analysis task
        tasks.append(Task(
            description=f"""Analyze the following tasks and past experience matches:
                          Tasks: {rfp_tasks}
                          Past Experience: {past_experience}
                          
                          Create a detailed analysis of how our past experience
                          demonstrates capability for each task.""",
            agent=analyzer
        ))

        # Writing task
        tasks.append(Task(
            description="""Using the analysis, create proposal sections that:
                          1. Address each RFP task
                          2. Incorporate specific past performance examples
                          3. Demonstrate clear capability through evidence""",
            agent=writer
        ))

        # Editing task
        tasks.append(Task(
            description="""Review and enhance the proposal to:
                          1. Strengthen past performance references
                          2. Ensure clear alignment with requirements
                          3. Polish language and presentation""",
            agent=editor
        ))

        return tasks

def main():
    st.set_page_config(page_title="AI Proposal Writer", layout="wide")
    st.title("🚀 Advanced AI Proposal Writer")
    st.markdown("---")

    api_key = st.text_input("OpenAI API Key", type="password")

    if not api_key:
        st.error("Please provide your OpenAI API Key.")
        return

    col1, col2 = st.columns(2)
    
    with col1:
        rfp_file = st.file_uploader("Upload RFP Document (.docx)", type=['docx'])
    
    with col2:
        past_projects_files = st.file_uploader(
            "Upload Past Project Documents (.docx)", 
            type=['docx'],
            accept_multiple_files=True
        )

    if rfp_file and past_projects_files and st.button("Generate Proposal"):
        try:
            proposal_writer = ProposalWriter(api_key)
            process_output_expander = st.expander("Processing Output:", expanded=True)
            sys.stdout = StreamToExpander(process_output_expander)

            with st.spinner("Processing documents..."):
                # Process RFP
                rfp_content = proposal_writer.doc_processor.extract_text_from_docx(rfp_file.read())
                rfp_tasks = proposal_writer.doc_processor.extract_tasks(rfp_content)

                # Process past projects
                past_projects = []
                for file in past_projects_files:
                    content = proposal_writer.doc_processor.extract_text_from_docx(file.read())
                    past_projects.append({
                        'title': file.name,
                        'content': content
                    })

                # Create vector store from past projects
                vector_store = proposal_writer.doc_processor.create_document_embeddings(past_projects)

                # Find relevant experience for each task
                past_experience = {}
                for task in rfp_tasks:
                    relevant_exp = proposal_writer.doc_processor.find_relevant_experience(task, vector_store)
                    summary = proposal_writer.doc_processor.summarize_experience(relevant_exp)
                    past_experience[task] = {
                        'examples': relevant_exp,
                        'summary': summary
                    }

                # Generate proposal
                analyzer, writer, editor = proposal_writer.setup_agents()
                tasks = proposal_writer.create_tasks(
                    analyzer, writer, editor, rfp_tasks, past_experience
                )

                crew = Crew(
                    agents=[analyzer, writer, editor],
                    tasks=tasks,
                    process=Process.sequential,
                    verbose=True
                )
                result = crew.kickoff()

                # Display results
                st.markdown("### Generated Proposal")
                st.text_area("Proposal Content", result, height=400)

                # Save and provide download
                filename = "Enhanced_Proposal.txt"
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(str(result))
                
                with open(filename, "rb") as f:
                    st.download_button(
                        label="Download Proposal",
                        data=f,
                        file_name=filename,
                        mime="text/plain"
                    )
                if os.path.exists(filename):
                    os.remove(filename)

        except Exception as e:
            st.error(f"An error occurred: {e}")
            st.error("Please try again or contact support if the issue persists.")

if __name__ == "__main__":
    main()
    